from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


AUTHOR = "Mateus Diniz Gottardi"
SUBJECT = "Sistema de Recomendação Baseado em Ontologias Fuzzy para Comércio Eletrônico"
ITALIC_PATTERN = re.compile(r"(\*[^*]+\*)")
ALPHA_ITEM_PATTERN = re.compile(r"^[a-z]\)\s+(.+)$")


def _set_run_font(
    run,
    *,
    size: int,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

    properties = run._element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Times New Roman")

    language = properties.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        properties.append(language)
    language.set(qn("w:val"), "pt-BR")


def _set_style_font(style, *, size: int, bold: bool = False) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)

    properties = style.element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Times New Roman")


def _add_inline_runs(paragraph, text: str, *, size: int, bold: bool = False) -> None:
    for part in ITALIC_PATTERN.split(text):
        if not part:
            continue
        italic = part.startswith("*") and part.endswith("*")
        content = part[1:-1] if italic else part
        run = paragraph.add_run(content)
        _set_run_font(run, size=size, bold=bold, italic=italic)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _set_style_font(normal, size=12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    heading_1 = document.styles["Heading 1"]
    _set_style_font(heading_1, size=12, bold=True)
    heading_1.paragraph_format.line_spacing = 1.5
    heading_1.paragraph_format.space_before = Pt(0)
    heading_1.paragraph_format.space_after = Pt(18)
    heading_1.paragraph_format.keep_with_next = True
    heading_1.paragraph_format.page_break_before = True

    heading_2 = document.styles["Heading 2"]
    _set_style_font(heading_2, size=11, bold=True)
    heading_2.paragraph_format.line_spacing = 1.5
    heading_2.paragraph_format.space_before = Pt(18)
    heading_2.paragraph_format.space_after = Pt(6)
    heading_2.paragraph_format.keep_with_next = True


def _next_numbering_id(elements, attribute: str) -> int:
    values = [int(element.get(qn(attribute))) for element in elements]
    return max(values, default=0) + 1


def _create_alpha_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_elements = numbering.findall(qn("w:abstractNum"))
    num_elements = numbering.findall(qn("w:num"))
    abstract_id = _next_numbering_id(abstract_elements, "w:abstractNumId")
    num_id = _next_numbering_id(num_elements, "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "lowerLetter")
    level.append(number_format)

    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1)")
    level.append(level_text)

    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)

    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    paragraph_properties.append(tabs)

    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "720")
    indentation.set(qn("w:hanging"), "360")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    num.append(abstract_reference)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    number_properties.append(level)
    number_properties.append(number)
    properties.insert(0, number_properties)


def _add_heading(document: Document, text: str, *, level: int) -> None:
    style_name = "Heading 1" if level == 1 else "Heading 2"
    paragraph = document.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_inline_runs(
        paragraph,
        text.upper(),
        size=12 if level == 1 else 11,
        bold=True,
    )


def _add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    _add_inline_runs(paragraph, text, size=12)


def _add_alpha_item(document: Document, text: str, *, num_id: int) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _apply_numbering(paragraph, num_id)
    _add_inline_runs(paragraph, text, size=12)


def export_markdown_chapter(
    input_path: Path,
    output_path: Path,
    *,
    document_title: str,
) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(2)
    section.footer_distance = Cm(1.25)

    _configure_styles(document)
    alpha_num_id = _create_alpha_numbering(document)
    document.core_properties.title = document_title
    document.core_properties.subject = SUBJECT
    document.core_properties.author = AUTHOR

    markdown_text = input_path.read_text(encoding="utf-8")
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            _add_heading(document, line[2:].strip(), level=1)
        elif line.startswith("## "):
            _add_heading(document, line[3:].strip(), level=2)
        else:
            alpha_match = ALPHA_ITEM_PATTERN.match(line)
            if alpha_match:
                _add_alpha_item(document, alpha_match.group(1), num_id=alpha_num_id)
            else:
                _add_body_paragraph(document, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
