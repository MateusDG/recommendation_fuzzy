from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


INTRODUCAO_MD = """# 1 INTRODUÇÃO

## 1.1 Contextualização do problema de pesquisa

Sistemas de recomendação ocupam posição estratégica em plataformas de comércio eletrônico porque reduzem o custo cognitivo da escolha. O usuário raramente percorre todo o catálogo disponível; ele navega por uma fração limitada de produtos, compara poucos atributos e decide a partir de sinais incompletos. Para a loja, a recomendação atua como mecanismo de organização do catálogo, exposição de itens relevantes e aumento de aderência entre intenção de compra e oferta. O problema deixa de ser apenas apresentar produtos populares. A tarefa técnica passa a ser interpretar preferências, inferir contexto e ordenar alternativas sob incerteza, aspecto recorrente na literatura de sistemas de recomendação (Adomavicius; Tuzhilin, 2005; Isinkaye; Folajimi; Ojokoh, 2015).

Modelos tradicionais de recomendação, especialmente os baseados em filtragem colaborativa, dependem de padrões recorrentes de interação entre usuários e itens. Essa dependência produz bons resultados quando há grande volume de avaliações, compras, cliques ou visualizações. O desempenho se fragiliza quando a matriz usuário-item é esparsa. Esse ponto é crítico em mercados de nicho, como comércios eletrônicos de produtos residenciais de alto padrão. Um catálogo de alto padrão, semelhante ao domínio de atuação da Kouzina Club, tende a envolver itens de maior valor financeiro, ciclos de compra mais longos e menor repetição transacional. O usuário não compra uma adega climatizada, uma cafeteira de alto padrão ou um conjunto de metais de banheiro com a mesma frequência com que compra itens de consumo recorrente.

A esparsidade afeta tanto usuários quanto produtos. Usuários novos ou pouco ativos não oferecem histórico suficiente para um perfil colaborativo robusto. Produtos novos, caros ou especializados podem permanecer com poucas avaliações por longos períodos. Esse é o problema de cold start, discutido de forma específica por Yuan e Hernandez (2023). Em plataformas varejistas de amplo alcance, a popularidade pode compensar parte dessa limitação. Em um mercado especializado, porém, recomendar apenas os itens mais avaliados pode distorcer a experiência, pois produtos de alto valor agregado nem sempre possuem grande volume de interações. A ausência de dados não significa ausência de relevância.

Outro problema aparece na natureza da preferência. A intenção de compra em produtos de alto padrão não é binária. Um usuário pode ter alta afinidade por cozinha gourmet, preferência moderada por marcas reconhecidas, baixa sensibilidade a preço e interesse parcial por materiais como aço inox ou acabamento fosco. Essas dimensões não se encaixam bem em classificações rígidas do tipo gosta ou não gosta. A teoria dos conjuntos fuzzy oferece uma forma mais adequada de representar esse comportamento, pois permite atribuir graus de pertinência a conceitos linguísticos como preço alto, avaliação excelente, produto de alto padrão, popularidade média ou compatibilidade elevada (Jain; Gupta, 2018).

Ontologias de domínio complementam essa representação. Um sistema baseado apenas em ratings observa interações, mas não compreende a estrutura semântica do catálogo. Já uma ontologia permite explicitar relações entre produto, categoria, ambiente, material, marca, função e perfil de usuário. Uma cafeteira espresso, por exemplo, pode ser modelada como produto de cozinha e área gourmet, associada a eletrodoméstico compacto, uso doméstico de alto padrão, faixa de preço elevada e atributos de design. Essa estrutura cria uma base para inferir recomendações mesmo quando o histórico de interação é limitado, aproximando a recomendação de abordagens baseadas em conhecimento e modelagem semântica (Guo et al., 2022; Guia; Silva; Bernardino, 2019).

O presente trabalho parte dessa interseção: recomendação personalizada, lógica fuzzy e ontologias aplicadas a comércio eletrônico. Trabalhos como os de Karthik e Ganapathy (2021) e Guia, Silva e Bernardino (2019) indicam que a combinação entre conhecimento de domínio, ontologias e técnicas de recomendação pode apoiar sistemas mais interpretáveis em comércio eletrônico. O recorte empírico adota o dataset Amazon Reviews 2023, no subconjunto Home and Kitchen, por sua proximidade com o domínio de produtos residenciais, cozinha, área gourmet e banheiro (Hou et al., 2024). A escolha não pressupõe que o dataset classifique oficialmente itens como luxo. O conceito de produto de alto padrão será tratado como uma construção operacional, derivada de atributos observáveis como preço relativo, avaliação média, número de avaliações, marca, categoria, materiais e termos associados a acabamento, design ou qualidade percebida.

## 1.2 Justificativa e Motivação

A justificativa deste trabalho está na limitação prática dos recomendadores tradicionais quando aplicados a catálogos especializados. A filtragem colaborativa funciona melhor quando há massa crítica de interações. No comércio eletrônico de alto padrão, essa massa crítica costuma ser irregular. Produtos caros recebem menos avaliações, usuários compram com menor frequência e a comparação entre itens envolve atributos subjetivos que não aparecem diretamente na matriz de ratings. O resultado é um risco técnico claro: o sistema pode privilegiar popularidade genérica em vez de compatibilidade real com o perfil do usuário. A literatura clássica já aponta que sistemas de recomendação precisam combinar diferentes fontes de informação quando uma única técnica não representa adequadamente o problema (Adomavicius; Tuzhilin, 2005; Burke, 2002).

Esse risco é mais relevante quando o produto recomendado possui alto valor financeiro. Uma recomendação inadequada de item simples pode gerar apenas baixo engajamento. Uma recomendação inadequada de produto de alto padrão pode comprometer a percepção de qualidade da plataforma. Nesse tipo de mercado, confiança e explicabilidade importam. O usuário precisa reconhecer algum sentido na recomendação: categoria compatível, ambiente de uso coerente, faixa de preço aceitável, avaliação positiva, material desejado ou proximidade com produtos anteriormente demonstrados como relevantes. A explicabilidade em sistemas de recomendação é tratada como fator relevante para transparência e aceitação das sugestões (Zhang; Chen, 2020).

A lógica fuzzy contribui justamente nesse ponto. Ela permite modelar a gradualidade da intenção de compra sem forçar fronteiras artificiais. Um produto pode ser parcialmente de alto padrão, não apenas de alto padrão ou fora desse perfil. Um usuário pode ter média afinidade por uma categoria e alta afinidade por determinado ambiente. A recomendação pode combinar esses graus e produzir um escore de compatibilidade mais expressivo do que uma regra rígida. A abordagem também permite controlar situações de risco. Se o produto possui preço elevado e a confiança no perfil do usuário é baixa, a recomendação pode receber prioridade moderada, ainda que sua avaliação média seja alta. Esse tipo de raciocínio é coerente com o uso da lógica fuzzy em sistemas de recomendação (Jain; Gupta, 2018; Karthik; Ganapathy, 2021).

A ontologia agrega uma segunda camada de valor. Ela organiza o conhecimento de domínio e permite que o sistema utilize relações semânticas para compensar a falta de interações. Em cenários de cold start, essa característica é decisiva. Um usuário com apenas uma avaliação positiva em Coffee, Tea & Espresso pode ainda receber recomendações relacionadas a cozinha gourmet, eletrodomésticos de alto padrão ou acessórios compatíveis, desde que a proximidade semântica esteja representada. O sistema deixa de depender exclusivamente de usuários semelhantes e passa a usar conhecimento estruturado sobre produtos, princípio associado a sistemas baseados em conhecimento e recomendação semântica (Guo et al., 2022; Guia; Silva; Bernardino, 2019).

A escolha do Amazon Reviews 2023 - Home and Kitchen se justifica por oferecer uma base pública e ampla para simular um ambiente de varejo digital. A validação inicial do schema confirmou categorias próximas ao domínio de interesse, como Kitchen & Dining, Dining & Entertaining, Kitchen Utensils & Gadgets, Storage & Organization, Bath, Bathroom Accessories, Dinnerware & Serveware e Glassware & Drinkware, além de campos úteis como título, descrição, características, preço, avaliação média e número de avaliações. Esse recorte permite construir uma ponte metodológica entre uma base pública e o contexto de uma empresa de nicho, como a Kouzina Club, sem depender de dados proprietários sensíveis. O uso de dados públicos também favorece reprodutibilidade experimental, desde que o recorte, os filtros e os parâmetros sejam documentados (Hou et al., 2024).

A motivação acadêmica também é direta. A literatura sobre sistemas de recomendação registra limitações recorrentes associadas a cold start, esparsidade, baixa interpretabilidade e dificuldade de representar preferências graduais (Yuan; Hernandez, 2023; Zhang; Chen, 2020). Trabalhos sobre ontologias fuzzy e recomendação baseada em conhecimento indicam caminhos promissores, mas ainda há espaço para investigar uma solução aplicada a comércio eletrônico especializado em produtos de alto padrão para casa, cozinha, área gourmet e banheiro (Karthik; Ganapathy, 2021; Guia; Silva; Bernardino, 2019). O trabalho busca ocupar esse espaço com um protótipo avaliável, documentado e comparável a baselines simples.

## 1.3 Objetivos (Geral e Específicos)

O objetivo geral deste trabalho é desenvolver e avaliar um sistema de recomendação híbrido baseado em ontologias fuzzy, capaz de modelar preferências graduais de usuários e gerar recomendações personalizadas em um contexto de comércio eletrônico voltado a produtos residenciais de alto padrão. A caracterização do sistema como híbrido decorre da combinação entre conhecimento de domínio, regras fuzzy, perfis de usuário e técnicas de ranking, alinhando-se à noção de integração de métodos em sistemas recomendadores (Burke, 2002).

Para alcançar esse objetivo, são definidos os seguintes objetivos específicos:

a) investigar os fundamentos de sistemas de recomendação, filtragem colaborativa, recomendação baseada em conteúdo, lógica fuzzy, ontologias de domínio e estratégias para mitigação de cold start (Adomavicius; Tuzhilin, 2005; Jain; Gupta, 2018; Yuan; Hernandez, 2023);

b) selecionar e preparar um recorte do dataset Amazon Reviews 2023 - Home and Kitchen compatível com produtos de cozinha, área gourmet, eletrodomésticos, utensílios de alto padrão, decoração funcional e banheiro (Hou et al., 2024);

c) definir uma representação operacional para produto de alto padrão/luxo, baseada em critérios observáveis como preço relativo, avaliação, popularidade, marca, categoria, materiais e atributos textuais;

d) modelar uma ontologia inicial de produtos, contemplando entidades como usuário, produto, categoria, ambiente, marca, faixa de preço, avaliação, atributo, material, estilo, função, preferência e recomendação (Guia; Silva; Bernardino, 2019);

e) construir perfis fuzzy de usuários a partir de interações positivas, negativas ou neutras, considerando afinidade por categoria, ambiente, marca, faixa de preço, popularidade e compatibilidade com produtos de alto padrão (Karthik; Ganapathy, 2021);

f) implementar um motor de inferência fuzzy para calcular escores de compatibilidade entre usuários e produtos candidatos (Jain; Gupta, 2018);

g) gerar rankings Top-N acompanhados de explicações textuais, indicando os fatores que contribuíram para cada recomendação (Zhang; Chen, 2020);

h) comparar o modelo proposto com baselines como popularidade, recomendação por categoria preferida, abordagem baseada em conteúdo simples e, se o escopo permitir, filtragem colaborativa (Isinkaye; Folajimi; Ojokoh, 2015);

i) avaliar os resultados por métricas adequadas a recomendação Top-N, como Precision@K, Recall@K, NDCG@K, Hit Rate@K e Coverage, incluindo cenários gerais e cenários de cold start (Yuan; Hernandez, 2023).

## 1.4 Organização da monografia

Esta monografia está organizada em sete capítulos. O primeiro capítulo apresenta a introdução, com a contextualização do problema, a justificativa, a motivação, os objetivos e a estrutura geral do trabalho.

O segundo capítulo aborda a fundamentação teórica. Serão discutidos os conceitos de sistemas de recomendação em comércio eletrônico, filtragem colaborativa, recomendação baseada em conteúdo, lógica fuzzy, teoria dos conjuntos fuzzy, ontologias, representação semântica e arquiteturas de aplicação. Essa base conceitual sustenta a escolha por uma abordagem híbrida, interpretável e orientada por conhecimento de domínio.

O terceiro capítulo apresenta os trabalhos relacionados. A análise considera estudos sobre recomendação fuzzy, ontologias aplicadas a perfis de usuário, sistemas de recomendação semânticos, estratégias para cold start e soluções voltadas a comércio eletrônico. A comparação busca identificar lacunas técnicas e metodológicas que justifiquem o modelo proposto.

O quarto capítulo descreve a metodologia. Nele são detalhados o levantamento e a preparação dos dados, o uso do Amazon Reviews 2023 - Home and Kitchen, os critérios de recorte do domínio, a construção da ontologia fuzzy, a definição das variáveis linguísticas, as regras de inferência, a arquitetura do protótipo e a estratégia de extração ou simulação de preferências a partir das interações disponíveis no dataset.

O quinto capítulo trata do desenvolvimento do protótipo. São descritos o motor de recomendação, a integração entre ontologia e regras fuzzy, os componentes de processamento, os mecanismos usados para gerar rankings Top-N e as explicações associadas às recomendações. O foco desse capítulo é mostrar como a proposta conceitual foi transformada em artefato computacional avaliável.

O sexto capítulo apresenta os testes e a validação experimental. Serão descritas as métricas utilizadas, os baselines de comparação, os cenários de cold start, os resultados obtidos e a análise crítica do desempenho do sistema. A avaliação deve verificar não apenas a precisão das recomendações, mas também cobertura, interpretabilidade e comportamento sob baixa disponibilidade de dados.

O sétimo capítulo reúne as conclusões e os trabalhos futuros. Essa parte retoma os objetivos, discute os resultados alcançados, aponta limitações do dataset e da modelagem adotada, e indica possíveis extensões, como uso de RDF/OWL, enriquecimento semântico por embeddings, integração com dados reais de navegação e refinamento das funções de pertinência fuzzy.
"""


INPUT_PATH = Path("outputs") / "Capitulo_1_Introducao.md"
OUTPUT_PATH = Path("outputs") / "Capitulo_1_Introducao.docx"


def set_run_font(run, *, bold: bool = False, size: int = 12) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_spacing(paragraph, *, first_line: bool = True) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(6)
    if first_line:
        fmt.first_line_indent = Cm(1.25)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, first_line=False)
    run = paragraph.add_run(text)
    set_run_font(run, bold=True, size=12)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    is_objective_item = bool(re.match(r"^[a-i]\)", text))
    set_paragraph_spacing(paragraph, first_line=not is_objective_item)
    if is_objective_item:
        paragraph.paragraph_format.left_indent = Cm(1.25)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def export_docx(markdown_text: str, output_path: Path) -> None:
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            add_heading(document, line[2:].strip(), level=1)
        elif line.startswith("## "):
            add_heading(document, line[3:].strip(), level=2)
        else:
            add_body_paragraph(document, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    export_docx(INPUT_PATH.read_text(encoding="utf-8"), OUTPUT_PATH)
    print(f"Arquivo gerado: {OUTPUT_PATH}")
